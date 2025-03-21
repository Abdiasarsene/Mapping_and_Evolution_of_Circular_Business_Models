import pandas as pd
from docx import Document

def csv_to_word(csv_file, word_file):
    # Lire le CSV
    df = pd.read_csv(csv_file)

    # Créer un document Word
    doc = Document()
    
    # Ajouter un titre
    doc.add_heading('Données du fichier CSV', level=1)

    # Ajouter un tableau avec le nombre de lignes et colonnes du CSV
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Ajouter les en-têtes
    for j, column in enumerate(df.columns):
        table.cell(0, j).text = column

    # Remplir les cellules avec les données
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)

    # Sauvegarder le document Word
    doc.save(word_file)
    print(f"Fichier Word généré : {word_file}")

# Chemin de ton fichier CSV
csv_path = r"D:\Projects\IT\Data Science & IA\Mapping_Evolution_of_Circular_Business_Models\Text_data\circular_economic.csv"
word_path = r"D:\Projects\IT\Data Science & IA\Mapping_Evolution_of_Circular_Business_Models\Text_data\circular_economic.docx"

# Exécution de la conversion
csv_to_word(csv_path, word_path)
